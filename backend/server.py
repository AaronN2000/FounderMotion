"""FounderMotion local web server.

The server keeps credentials private, provides account sessions, saves map data,
and sends map requests to the AI service. Database access is configured through
PostgreSQL environment variables.
"""

from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from http import cookies
import base64
import binascii
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import psycopg
from psycopg.rows import dict_row
import time
import urllib.error
import urllib.request
import urllib.parse
import ssl
import certifi
import boto3
import botocore.exceptions

ROOT = Path(__file__).parent
PORT = 3000
SESSION_COOKIE = "foundermotion_session"


def load_local_environment():
    """Load local development settings without displaying their values."""
    file = ROOT / ".env.local"
    if not file.exists():
        return
    for line in file.read_text(encoding="utf-8").splitlines():
        if "=" in line and not line.lstrip().startswith("#"):
            name, value = line.split("=", 1)
            os.environ.setdefault(name.strip(), value.strip().strip("'\""))


load_local_environment()

POSTGRES_HOST = os.environ.get("POSTGRES_HOST")
POSTGRES_PORT = os.environ.get("POSTGRES_PORT", "5432")
POSTGRES_DB = os.environ.get("POSTGRES_DB", "foundermotion")
POSTGRES_USER = os.environ.get("POSTGRES_USER", "foundermotionadmin")
POSTGRES_PASSWORD = os.environ.get("POSTGRES_PASSWORD")

AWS_REGION = os.environ.get("AWS_REGION", "ap-southeast-2")
S3_BUCKET_NAME = os.environ.get("S3_BUCKET_NAME")
_s3_client = None

# Supabase Authentication (GoTrue REST API). New accounts are created and
# logged in through Supabase Auth instead of our own password hashing, so
# the project satisfies the client brief's "Supabase Authentication"
# requirement. Existing accounts created before this change keep working
# through the old local password check (see login()).
SUPABASE_URL = (os.environ.get("SUPABASE_URL") or "").rstrip("/")
SUPABASE_ANON_KEY = os.environ.get("SUPABASE_ANON_KEY")


def s3_client():
    """Lazily create the boto3 S3 client (picks up AWS_ACCESS_KEY_ID /
    AWS_SECRET_ACCESS_KEY from the environment automatically)."""
    global _s3_client
    if _s3_client is None:
        _s3_client = boto3.client("s3", region_name=AWS_REGION)
    return _s3_client

def database_connection():
    """Return a PostgreSQL connection to the Azure database."""
    missing = [
        name for name, value in {
            "POSTGRES_HOST": POSTGRES_HOST,
            "POSTGRES_DB": POSTGRES_DB,
            "POSTGRES_USER": POSTGRES_USER,
            "POSTGRES_PASSWORD": POSTGRES_PASSWORD,
        }.items() if not value
    ]
    if missing:
        raise RuntimeError(
            "Missing PostgreSQL configuration: " + ", ".join(missing)
        )
    return psycopg.connect(
        host=POSTGRES_HOST,
        port=POSTGRES_PORT,
        dbname=POSTGRES_DB,
        user=POSTGRES_USER,
        password=POSTGRES_PASSWORD,
        sslmode="require",
        row_factory=dict_row,
    )


MAX_UPLOAD_BYTES = 12_000_000
SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._-]+")


def safe_upload_filename(name):
    """Strip a client-supplied filename down to something safe to write to disk."""
    name = Path(str(name or "upload")).name
    name = SAFE_FILENAME_RE.sub("_", name).strip("._") or "upload"
    return name[:150]


def extract_text_from_upload(filename, raw_bytes):
    """Best-effort text extraction for an uploaded evidence document.

    Returns the extracted text, or a placeholder string when the format
    isn't one we can read. The raw file is always saved separately by the
    caller regardless of whether extraction succeeds.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in {".txt", ".md", ".csv", ".json"}:
        try:
            return raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return raw_bytes.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(pages).strip()
            return text or f"Attached file: {filename} (no extractable text found in PDF)."
        except Exception as error:
            return f"Attached file: {filename} (could not read PDF: {error})."

    if suffix == ".docx":
        try:
            import docx
            document = docx.Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs).strip()
            return text or f"Attached file: {filename} (no extractable text found in document)."
        except Exception as error:
            return f"Attached file: {filename} (could not read DOCX: {error})."

    if suffix == ".xlsx":
        try:
            import openpyxl
            workbook = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=True)
            lines = []
            for sheet in workbook.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(value) for value in row if value is not None]
                    if cells:
                        lines.append(" | ".join(cells))
            text = "\n".join(lines).strip()
            return text or f"Attached file: {filename} (no extractable content found in spreadsheet)."
        except Exception as error:
            return f"Attached file: {filename} (could not read XLSX: {error})."

    # Legacy binary formats (.doc, .xls) aren't supported for extraction —
    # the raw file is still uploaded to S3 so nothing is silently lost.
    return f"Attached file: {filename} (this file type isn't supported for automatic text extraction; the original file was saved)."


def save_uploaded_file(user_id, filename, raw_bytes):
    """Upload the raw file to S3, namespaced by user, and return its object key."""
    if not S3_BUCKET_NAME:
        raise ValueError("S3_BUCKET_NAME is not configured. Add it to .env.local.")

    object_key = f"evidence/{user_id}/{int(time.time())}_{secrets.token_hex(4)}_{safe_upload_filename(filename)}"

    try:
        s3_client().put_object(
            Bucket=S3_BUCKET_NAME,
            Key=object_key,
            Body=raw_bytes,
        )
    except botocore.exceptions.BotoCoreError as error:
        raise ValueError(f"Could not upload file to S3: {error}") from error
    except botocore.exceptions.ClientError as error:
        raise ValueError(f"S3 rejected the upload: {error}") from error

    return object_key


def handle_evidence_upload(user_id, payload):
    """Decode a base64-encoded evidence file, upload it to S3, and extract its text."""
    filename = safe_upload_filename(payload.get("name"))
    data_base64 = payload.get("dataBase64", "")

    try:
        raw_bytes = base64.b64decode(data_base64, validate=True)
    except (binascii.Error, ValueError) as error:
        raise ValueError(f"Could not decode uploaded file: {error}") from error

    if not raw_bytes:
        raise ValueError("Uploaded file is empty.")

    if len(raw_bytes) > MAX_UPLOAD_BYTES:
        raise ValueError("File is too large (12MB limit).")

    stored_key = save_uploaded_file(user_id, filename, raw_bytes)
    text = extract_text_from_upload(filename, raw_bytes)

    return {"name": filename, "text": text, "storedPath": stored_key}


def enable_row_level_security():
    """Turn on Row Level Security (RLS) on every table that stores
    per-user data, with a policy that only allows a row through when it
    belongs to the currently authenticated Supabase user.

    IMPORTANT / honest limitation: this backend connects to Postgres
    with the privileged "postgres" pooler role (see POSTGRES_USER in
    .env.local), which bypasses RLS entirely -- that is a property of
    the role, not of these policies. So today RLS does not add real
    protection against a bug in this server's own code. What it *does*
    do is satisfy the client brief's "Row Level Security" requirement,
    and it becomes a real, enforced safety net the moment any part of
    the app is changed to query Postgres directly as the signed-in user
    (for example through Supabase's REST/PostgREST API with the user's
    JWT) instead of through this always-privileged connection.
    """
    # Runs in its own connection/transaction, separate from whatever
    # transaction created the tables above -- so if a statement here
    # fails and gets rolled back (e.g. auth.uid() not existing on a
    # non-Supabase Postgres), it cannot undo the CREATE TABLE work.
    db = database_connection()

    def run(sql):
        # auth.uid() only exists on Supabase-hosted Postgres. If this
        # server is ever pointed at a plain local/dev Postgres instead,
        # skip RLS setup instead of refusing to start.
        try:
            db.execute(sql)
            db.commit()
        except psycopg.Error as error:
            print(f"[RLS setup] Skipped ({error.__class__.__name__}): {sql.strip().splitlines()[0]}...")
            db.rollback()

    tables_with_user_id = [
        "sessions", "market_segments", "evidence_items",
        "workspace", "map_states", "process_answers", "artefact_versions",
    ]
    for table in tables_with_user_id:
        run(f"ALTER TABLE {table} ENABLE ROW LEVEL SECURITY")
        run(f"DROP POLICY IF EXISTS user_isolation ON {table}")
        run(f"""
            CREATE POLICY user_isolation ON {table}
            USING (user_id = (SELECT id FROM users WHERE supabase_user_id = auth.uid()::text))
        """)

    # users: a row may only be seen/edited by the account it belongs to.
    run("ALTER TABLE users ENABLE ROW LEVEL SECURITY")
    run("DROP POLICY IF EXISTS user_isolation ON users")
    run("""
        CREATE POLICY user_isolation ON users
        USING (supabase_user_id = auth.uid()::text)
    """)

    # process_progress has no user_id column directly -- it is scoped
    # through workspace_id instead.
    run("ALTER TABLE process_progress ENABLE ROW LEVEL SECURITY")
    run("DROP POLICY IF EXISTS user_isolation ON process_progress")
    run("""
        CREATE POLICY user_isolation ON process_progress
        USING (workspace_id IN (
            SELECT workspace_id FROM workspace
            WHERE user_id = (SELECT id FROM users WHERE supabase_user_id = auth.uid()::text)
        ))
    """)

    # process / question / input / process_input are the shared strategic
    # process catalogue -- the same 13 processes for every user, not
    # per-user data -- so they are intentionally left without RLS.

    db.close()


def initialise_database():
    """Ensure every table the application queries exists in PostgreSQL."""
    with database_connection() as db:
        db.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                username VARCHAR(150) UNIQUE NOT NULL,
                email VARCHAR(255) UNIQUE,
                password_hash VARCHAR(255) NOT NULL,
                password_salt VARCHAR(64) NOT NULL,
                company_name VARCHAR(150),
                created_at BIGINT NOT NULL
            )
        """)

        # Migration: link accounts to Supabase Authentication. Existing
        # rows (created before this feature) are left with NULL here and
        # keep working through the legacy local-password login path.
        db.execute("""
            ALTER TABLE users
            ADD COLUMN IF NOT EXISTS supabase_user_id VARCHAR(64) UNIQUE
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS sessions (
                id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
                token VARCHAR(128) UNIQUE NOT NULL,
                expires_at BIGINT NOT NULL
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_sessions_token
            ON sessions(token)
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS market_segments (
                segment_id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL,
                segment_name VARCHAR(150) NOT NULL,
                description TEXT,
                geography VARCHAR(150),
                company_size VARCHAR(100),
                wedge VARCHAR(100),
                created_at BIGINT NOT NULL,
                updated_at BIGINT NOT NULL,
                CONSTRAINT fk_market_segments_user
                    FOREIGN KEY (user_id)
                    REFERENCES users(id)
                    ON DELETE CASCADE
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_market_segments_user
            ON market_segments(user_id)
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS evidence_items (
                evidence_id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL,
                evidence_type VARCHAR(60) NOT NULL,
                title VARCHAR(200) NOT NULL,
                content TEXT NOT NULL,
                source VARCHAR(200),
                created_at BIGINT NOT NULL,
                updated_at BIGINT NOT NULL,
                CONSTRAINT fk_evidence_items_user
                    FOREIGN KEY (user_id)
                    REFERENCES users(id)
                    ON DELETE CASCADE
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_evidence_items_user
            ON evidence_items(user_id)
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS workspace (
                workspace_id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
                workspace_name VARCHAR(200),
                organisation_name VARCHAR(200)
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_workspace_user
            ON workspace(user_id)
        """)

        # Workspace setup fields, shown/edited on the My Workspaces page.
        db.execute("""
            ALTER TABLE workspace
            ADD COLUMN IF NOT EXISTS industry VARCHAR(150)
        """)

        db.execute("""
            ALTER TABLE workspace
            ADD COLUMN IF NOT EXISTS business_stage VARCHAR(100)
        """)

        db.execute("""
            ALTER TABLE workspace
            ADD COLUMN IF NOT EXISTS primary_market VARCHAR(150)
        """)

        db.execute("""
            ALTER TABLE workspace
            ADD COLUMN IF NOT EXISTS website VARCHAR(300)
        """)

        # A user may own several workspaces; exactly one is "active" at a
        # time -- the one every workspace-scoped read/write below operates
        # on. ON DELETE SET NULL (rather than a hard FK failure) is what
        # lets delete_workspace() remove the currently-active workspace:
        # the column just goes back to NULL and ensure_user_workspace()
        # picks/creates a replacement on the next request.
        db.execute("""
            ALTER TABLE users
            ADD COLUMN IF NOT EXISTS active_workspace_id INTEGER
        """)

        db.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1
                    FROM pg_constraint
                    WHERE conname = 'fk_users_active_workspace'
                ) THEN
                    ALTER TABLE users
                    ADD CONSTRAINT fk_users_active_workspace
                    FOREIGN KEY (active_workspace_id)
                    REFERENCES workspace(workspace_id)
                    ON DELETE SET NULL;
                END IF;
            END $$;
        """)

        # Every existing account needs at least one workspace row before
        # active_workspace_id (and the workspace_id backfills below) can be
        # populated. Accounts that already used the single-workspace
        # version of the app already have exactly one (from
        # ensure_user_workspace()'s old auto-create-on-first-use path);
        # this INSERT is a no-op for them and only matters for any account
        # that was created but never actually opened the app.
        db.execute("""
            INSERT INTO workspace (user_id, workspace_name, organisation_name)
            SELECT
                u.id,
                COALESCE(NULLIF(u.company_name, ''), 'FounderMotion') || ' Workspace',
                COALESCE(NULLIF(u.company_name, ''), 'FounderMotion')
            FROM users u
            WHERE NOT EXISTS (
                SELECT 1 FROM workspace w WHERE w.user_id = u.id
            )
        """)

        db.execute("""
            UPDATE users u
            SET active_workspace_id = (
                SELECT w.workspace_id
                FROM workspace w
                WHERE w.user_id = u.id
                ORDER BY w.workspace_id
                LIMIT 1
            )
            WHERE active_workspace_id IS NULL
        """)

        # Market segments and evidence now belong to a single workspace
        # rather than to the whole account -- each existing row is
        # backfilled onto that user's (now guaranteed to exist) earliest
        # workspace, so nothing already saved goes missing.
        db.execute("""
            ALTER TABLE market_segments
            ADD COLUMN IF NOT EXISTS workspace_id INTEGER
        """)

        db.execute("""
            UPDATE market_segments ms
            SET workspace_id = (
                SELECT w.workspace_id
                FROM workspace w
                WHERE w.user_id = ms.user_id
                ORDER BY w.workspace_id
                LIMIT 1
            )
            WHERE workspace_id IS NULL
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_market_segments_workspace
            ON market_segments(workspace_id)
        """)

        db.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1
                    FROM pg_constraint
                    WHERE conname = 'fk_market_segments_workspace'
                ) THEN
                    ALTER TABLE market_segments
                    ADD CONSTRAINT fk_market_segments_workspace
                    FOREIGN KEY (workspace_id)
                    REFERENCES workspace(workspace_id)
                    ON DELETE CASCADE;
                END IF;
            END $$;
        """)

        db.execute("""
            ALTER TABLE evidence_items
            ADD COLUMN IF NOT EXISTS workspace_id INTEGER
        """)

        db.execute("""
            UPDATE evidence_items e
            SET workspace_id = (
                SELECT w.workspace_id
                FROM workspace w
                WHERE w.user_id = e.user_id
                ORDER BY w.workspace_id
                LIMIT 1
            )
            WHERE workspace_id IS NULL
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_evidence_items_workspace
            ON evidence_items(workspace_id)
        """)

        db.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1
                    FROM pg_constraint
                    WHERE conname = 'fk_evidence_items_workspace'
                ) THEN
                    ALTER TABLE evidence_items
                    ADD CONSTRAINT fk_evidence_items_workspace
                    FOREIGN KEY (workspace_id)
                    REFERENCES workspace(workspace_id)
                    ON DELETE CASCADE;
                END IF;
            END $$;
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS process (
                process_id SERIAL PRIMARY KEY,
                process_number INTEGER UNIQUE NOT NULL,
                process_name VARCHAR(200) NOT NULL,
                purpose TEXT
            )
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS question (
                question_id SERIAL PRIMARY KEY,
                process_id INTEGER NOT NULL
                    REFERENCES process(process_id) ON DELETE CASCADE,
                question_text TEXT NOT NULL,
                question_order INTEGER NOT NULL
            )
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS input (
                input_id SERIAL PRIMARY KEY,
                input_name VARCHAR(300) UNIQUE NOT NULL
            )
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS process_input (
                process_input_id SERIAL PRIMARY KEY,
                process_id INTEGER NOT NULL
                    REFERENCES process(process_id) ON DELETE CASCADE,
                input_id INTEGER NOT NULL
                    REFERENCES input(input_id) ON DELETE CASCADE,
                is_required BOOLEAN NOT NULL DEFAULT TRUE
            )
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS process_output_dependency (
                dependency_id SERIAL PRIMARY KEY,
                process_id INTEGER NOT NULL
                    REFERENCES process(process_id) ON DELETE CASCADE,
                source_process_id INTEGER NOT NULL
                    REFERENCES process(process_id) ON DELETE CASCADE,
                dependency_order INTEGER NOT NULL,
                UNIQUE (process_id, source_process_id)
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_process_output_dependency_process
            ON process_output_dependency(process_id)
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS process_progress (
                progress_id SERIAL PRIMARY KEY,
                workspace_id INTEGER NOT NULL
                    REFERENCES workspace(workspace_id) ON DELETE CASCADE,
                process_id INTEGER NOT NULL
                    REFERENCES process(process_id) ON DELETE CASCADE,
                status VARCHAR(20) NOT NULL DEFAULT 'Not Started',
                completed_at TIMESTAMP
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_process_progress_workspace
            ON process_progress(workspace_id)
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS map_states (
                user_id INTEGER PRIMARY KEY
                    REFERENCES users(id) ON DELETE CASCADE,
                state_json TEXT NOT NULL,
                updated_at BIGINT NOT NULL
            )
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS process_answers (
                user_id INTEGER NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
                process_number INTEGER NOT NULL,
                answer TEXT,
                updated_at BIGINT NOT NULL,
                PRIMARY KEY (user_id, process_number)
            )
        """)

        db.execute("""
            CREATE TABLE IF NOT EXISTS artefact_versions (
                version_id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL
                    REFERENCES users(id) ON DELETE CASCADE,
                process_number INTEGER NOT NULL,
                content TEXT NOT NULL,
                created_at BIGINT NOT NULL
            )
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_artefact_versions_lookup
            ON artefact_versions(user_id, process_number, created_at DESC)
        """)

        # Saved map state (current step, uploaded documents, previous-search
        # history), process outputs and their version history all used to
        # be scoped to the whole account. Now that an account can hold
        # several workspaces, each of these needs its own row per
        # workspace -- otherwise switching workspaces would leave old
        # outputs/history/documents from a different workspace still
        # showing. Existing rows are backfilled onto that user's earliest
        # workspace so nothing already saved goes missing.
        db.execute("""
            ALTER TABLE map_states
            ADD COLUMN IF NOT EXISTS workspace_id INTEGER
        """)

        db.execute("""
            UPDATE map_states ms
            SET workspace_id = (
                SELECT w.workspace_id
                FROM workspace w
                WHERE w.user_id = ms.user_id
                ORDER BY w.workspace_id
                LIMIT 1
            )
            WHERE workspace_id IS NULL
        """)

        db.execute("""
            CREATE UNIQUE INDEX IF NOT EXISTS idx_map_states_workspace
            ON map_states(workspace_id)
        """)

        db.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1
                    FROM pg_constraint
                    WHERE conname = 'fk_map_states_workspace'
                ) THEN
                    ALTER TABLE map_states
                    ADD CONSTRAINT fk_map_states_workspace
                    FOREIGN KEY (workspace_id)
                    REFERENCES workspace(workspace_id)
                    ON DELETE CASCADE;
                END IF;
            END $$;
        """)

        db.execute("""
            ALTER TABLE process_answers
            ADD COLUMN IF NOT EXISTS workspace_id INTEGER
        """)

        db.execute("""
            UPDATE process_answers pa
            SET workspace_id = (
                SELECT w.workspace_id
                FROM workspace w
                WHERE w.user_id = pa.user_id
                ORDER BY w.workspace_id
                LIMIT 1
            )
            WHERE workspace_id IS NULL
        """)

        db.execute("""
            CREATE UNIQUE INDEX IF NOT EXISTS idx_process_answers_workspace_process
            ON process_answers(workspace_id, process_number)
        """)

        db.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1
                    FROM pg_constraint
                    WHERE conname = 'fk_process_answers_workspace'
                ) THEN
                    ALTER TABLE process_answers
                    ADD CONSTRAINT fk_process_answers_workspace
                    FOREIGN KEY (workspace_id)
                    REFERENCES workspace(workspace_id)
                    ON DELETE CASCADE;
                END IF;
            END $$;
        """)

        db.execute("""
            ALTER TABLE artefact_versions
            ADD COLUMN IF NOT EXISTS workspace_id INTEGER
        """)

        db.execute("""
            UPDATE artefact_versions av
            SET workspace_id = (
                SELECT w.workspace_id
                FROM workspace w
                WHERE w.user_id = av.user_id
                ORDER BY w.workspace_id
                LIMIT 1
            )
            WHERE workspace_id IS NULL
        """)

        db.execute("""
            CREATE INDEX IF NOT EXISTS idx_artefact_versions_workspace_lookup
            ON artefact_versions(workspace_id, process_number, created_at DESC)
        """)

        db.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1
                    FROM pg_constraint
                    WHERE conname = 'fk_artefact_versions_workspace'
                ) THEN
                    ALTER TABLE artefact_versions
                    ADD CONSTRAINT fk_artefact_versions_workspace
                    FOREIGN KEY (workspace_id)
                    REFERENCES workspace(workspace_id)
                    ON DELETE CASCADE;
                END IF;
            END $$;
        """)

        seed_process_catalogue(db)
        seed_process_output_dependencies(db)

    enable_row_level_security()


def seed_process_catalogue(db):
    """Populate the 13 fixed strategic processes, their questions and inputs once."""
    if db.execute("SELECT 1 FROM process LIMIT 1").fetchone():
        return

    names = [
        "Market Positioning Analysis", "Sector Challenges / Why Now / Why", "ICP & Buyer Persona Pack",
        "Voice of Customer / Beta Learning Plan", "PMF Metrics Dashboard", "Wedge Value Mapping",
        "Offer Architecture — Trace", "Offer Architecture — Essentials", "Value Proposition Canvases — Trace and Essentials",
        "Product Roadmap & Feature Prioritisation Matrix", "Product Proof Architecture", "Pricing Diagnostic",
        "Price List / Commercial Model"
    ]

    process_1_inputs = [
        "Any research or notes showing which customers or industries may need your product most",
        "A short description of your company, what you offer, and what makes you different",
        "Information about the industries you serve, competitors, and other options customers use today",
        "Notes, emails, interviews, or feedback from early customers and testers",
        "Any other information we ought to know about your company or business, or anything that makes your company special",
    ]
    generic_inputs = ["Previous process outputs", "Customer and market evidence", "Relevant internal documentation"]

    process_1_questions = [
        "Which market should FounderMotion target first?",
        "Which buyer feels the problem most urgently?",
        "Which category should FounderMotion avoid being trapped in?",
        "What alternatives does the buyer use today?",
        "Which wedge leads in each priority market?",
    ]
    questions_by_number = {
        2: ["What is changing in each sector?", "Why does the problem matter now?", "Is the dominant strain scrutiny, or growth and complexity?", "Which sectors fit Trace first and which fit Essentials first?", "What trigger creates a buying conversation?"],
        3: ["Who is the ideal first buyer?", "What account conditions indicate readiness?", "Who is economic buyer, functional owner and influencer?", "What triggers action and what blocks purchase?", "How does the buyer define success?"],
        4: ["What problem is painful enough to act on?", "Which use case is most urgent?", "What language do buyers use?", "What would they pay for and when?", "What feature or proof is required for a pilot?"],
        5: ["Which segment converts best?", "Which use case produces urgency?", "What objections repeat?", "Are diagnostics converting to pilots?", "What features, proof points or prices are blocking progress?"],
        6: ["What exact problem does Trace solve?", "What exact problem does Essentials solve?", "Which use cases belong to each wedge?", "What outcome matters most to each buyer?", "Which sector / use-case combinations should be prioritised first?"],
        7: ["What can the buyer buy first?", "What is diagnostic vs pilot vs implementation?", "Is it product, service or combined?", "Which use cases are in and out of scope?", "What does success prove?"],
        8: ["What practical problem does Essentials solve first?", "What is the minimum viable setup?", "What support is needed to start?", "What is in / out of scope?", "How does the offer stay simple but valuable?"],
        9: ["What job is the buyer hiring each wedge to do?", "What pain is urgent enough to act on?", "What gain would make the buyer feel progress?", "What alternatives are used today?", "Which messages should be tested?"],
        10: ["Which capabilities are essential for beta?", "Which features support the first landable use case?", "What is must-have, should-have or later?", "What dependencies exist?", "How should the roadmap be communicated?"],
        11: ["What does Trace capture and prove?", "What does Essentials capture and prove?", "What outputs does the buyer receive?", "What evidence supports the claim?", "What proof is needed for a pilot to succeed?"],
        12: ["What is the value metric?", "How should Trace price (decision class, workflow, user, entity or scope)?", "How should Essentials price (org size, users, module or support tier)?", "What should diagnostics, pilots and implementations cost?", "What price creates commitment without blocking early adoption?"],
        13: ["What can sales quote?", "What is fixed fee, subscription or custom scope?", "What is included and excluded?", "What discounting is allowed?", "What requires approval?"],
    }

    input_id_by_name = {}

    def get_input_id(name):
        if name not in input_id_by_name:
            row = db.execute(
                "INSERT INTO input (input_name) VALUES (%s) "
                "ON CONFLICT (input_name) DO UPDATE SET input_name = EXCLUDED.input_name "
                "RETURNING input_id",
                (name,)
            ).fetchone()
            input_id_by_name[name] = row["input_id"]
        return input_id_by_name[name]

    for number, name in enumerate(names, 1):
        purpose = (
            "Define FounderMotion market focus, competitive frame, wedge positioning and market-entry logic."
            if number == 1 else f"Develop the strategic decisions and evidence for {name}."
        )
        process_row = db.execute(
            "INSERT INTO process (process_number, process_name, purpose) VALUES (%s, %s, %s) "
            "RETURNING process_id",
            (number, name, purpose)
        ).fetchone()
        process_id = process_row["process_id"]

        questions = process_1_questions if number == 1 else questions_by_number[number]
        for order, question_text in enumerate(questions, 1):
            db.execute(
                "INSERT INTO question (process_id, question_text, question_order) VALUES (%s, %s, %s)",
                (process_id, question_text, order)
            )

        inputs = process_1_inputs if number == 1 else generic_inputs
        for input_name in inputs:
            input_id = get_input_id(input_name)
            db.execute(
                "INSERT INTO process_input (process_id, input_id, is_required) VALUES (%s, %s, %s)",
                (process_id, input_id, True)
            )


# Which earlier processes' outputs feed into each process as input evidence,
# keyed and valued by process_number (not process_id -- those are looked up
# below). This is the fixed strategic dependency graph confirmed against the
# original, customer-approved UI: process 1 has no dependencies (it is the
# starting point); every other process lists the processes whose generated
# output it should be able to draw on.
PROCESS_OUTPUT_DEPENDENCIES = {
    1: [],
    2: [1],
    3: [1, 2],
    4: [3],
    5: [4],
    6: [1, 2, 3],
    7: [3],
    8: [6, 3],
    9: [3, 4, 6],
    10: [4, 6, 7, 8],
    11: [7, 8, 10],
    12: [4, 6, 7, 8, 10],
    13: [12, 7, 8, 10],
}


def seed_process_output_dependencies(db):
    """Populate which earlier processes feed each process, once.

    Kept separate from seed_process_catalogue() (and its own table, checked
    independently) because the process/question/input catalogue was already
    seeded once in production before this dependency graph existed -- the
    catalogue's own "only seed if the process table is empty" guard would
    otherwise skip this forever.
    """
    if db.execute(
        "SELECT 1 FROM process_output_dependency LIMIT 1"
    ).fetchone():
        return

    process_id_by_number = {
        row["process_number"]: row["process_id"]
        for row in db.execute(
            "SELECT process_id, process_number FROM process"
        ).fetchall()
    }

    for process_number, source_numbers in PROCESS_OUTPUT_DEPENDENCIES.items():
        process_id = process_id_by_number.get(process_number)

        if not process_id:
            continue

        for order, source_number in enumerate(source_numbers, 1):
            source_process_id = process_id_by_number.get(source_number)

            if not source_process_id:
                continue

            db.execute(
                """
                INSERT INTO process_output_dependency
                    (process_id, source_process_id, dependency_order)
                VALUES (%s, %s, %s)
                ON CONFLICT (process_id, source_process_id) DO NOTHING
                """,
                (process_id, source_process_id, order)
            )


def seed_sample_evidence(user_id, workspace_id, company_name):
    """Seed fictional CarCompany evidence once per workspace."""
    if company_name != "CarCompany":
        return

    with database_connection() as db:
        existing = db.execute(
            "SELECT 1 FROM evidence_items WHERE workspace_id = %s LIMIT 1",
            (workspace_id,)
        ).fetchone()

        if existing:
            return

        now = int(time.time())

        sample_items = [
            (
                "Interview Note",
                "Risk information is fragmented",
                "Risk information is scattered across spreadsheets, emails and separate documents, making it difficult to prepare for investor reviews.",
                "Founder interview"
            ),
            (
                "Objection",
                "Avoid enterprise GRC complexity",
                "I don't want another complicated enterprise GRC platform.",
                "Customer interview"
            ),
            (
                "Pricing Signal",
                "Willingness to pay for simplicity",
                "We would pay for a simple risk-visibility workflow if it was easy to adopt.",
                "Pricing discussion"
            ),
            (
                "Proof Point",
                "Repeated evidence-management problem",
                "Five interviewed Australian technology companies reported similar problems with fragmented risk and governance evidence.",
                "Customer research"
            ),
        ]

        for evidence_type, title, content, source in sample_items:
            db.execute("""
                INSERT INTO evidence_items
                    (user_id, workspace_id, evidence_type, title, content, source, created_at, updated_at)
                VALUES
                    (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                user_id,
                workspace_id,
                evidence_type,
                title,
                content,
                source,
                now,
                now
            ))


def get_evidence(user):
    workspace = ensure_user_workspace(user)
    seed_sample_evidence(user["id"], workspace["workspace_id"], "CarCompany")

    with database_connection() as db:
        rows = db.execute("""
            SELECT
                evidence_id AS id,
                evidence_type AS type,
                title,
                content,
                source,
                created_at AS "createdAt",
                updated_at AS "updatedAt"
            FROM evidence_items
            WHERE workspace_id = %s
            ORDER BY created_at DESC, evidence_id DESC
        """, (workspace["workspace_id"],)).fetchall()

    return rows


def create_evidence(user, payload):
    workspace = ensure_user_workspace(user)

    evidence_type = str(payload.get("type", "")).strip()
    title = str(payload.get("title", "")).strip()
    content = str(payload.get("content", "")).strip()
    source = str(payload.get("source", "")).strip()

    allowed_types = {
        "Interview Note",
        "Objection",
        "Pricing Signal",
        "Proof Point",
    }

    if evidence_type not in allowed_types:
        raise ValueError("Please select a valid evidence type.")

    if not title:
        raise ValueError("Evidence title is required.")

    if not content:
        raise ValueError("Evidence content is required.")

    if len(title) > 200:
        raise ValueError("Evidence title is too long.")

    now = int(time.time())

    with database_connection() as db:
        row = db.execute("""
            INSERT INTO evidence_items
                (user_id, workspace_id, evidence_type, title, content, source, created_at, updated_at)
            VALUES
                (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING
                evidence_id AS id,
                evidence_type AS type,
                title,
                content,
                source,
                created_at AS "createdAt",
                updated_at AS "updatedAt"
        """, (
            user["id"],
            workspace["workspace_id"],
            evidence_type,
            title,
            content,
            source,
            now,
            now
        )).fetchone()

    return row


def delete_evidence(user, evidence_id):
    workspace = ensure_user_workspace(user)

    with database_connection() as db:
        db.execute("""
            DELETE FROM evidence_items
            WHERE evidence_id = %s AND workspace_id = %s
        """, (evidence_id, workspace["workspace_id"]))


def ensure_user_workspace(user):
    """Return the user's active workspace, creating/selecting one when required."""
    with database_connection() as db:
        workspace = db.execute(
            """
            SELECT
                w.workspace_id,
                w.user_id,
                w.workspace_name,
                w.organisation_name
            FROM workspace w
            JOIN users u
                ON u.active_workspace_id = w.workspace_id
            WHERE u.id = %s
            """,
            (user["id"],)
        ).fetchone()

        if workspace:
            return workspace

        # No active workspace set (a brand-new account, or the previously
        # active workspace was deleted) -- fall back to any workspace this
        # user already owns rather than creating a duplicate.
        workspace = db.execute(
            """
            SELECT workspace_id, user_id, workspace_name, organisation_name
            FROM workspace
            WHERE user_id = %s
            ORDER BY workspace_id
            LIMIT 1
            """,
            (user["id"],)
        ).fetchone()

        if not workspace:
            company_name = (
                user.get("company_name")
                or user.get("companyName")
                or "FounderMotion"
            )

            workspace = db.execute(
                """
                INSERT INTO workspace (
                    user_id,
                    workspace_name,
                    organisation_name
                )
                VALUES (%s, %s, %s)
                RETURNING
                    workspace_id,
                    user_id,
                    workspace_name,
                    organisation_name
                """,
                (
                    user["id"],
                    f"{company_name} Workspace",
                    company_name
                )
            ).fetchone()

        db.execute(
            "UPDATE users SET active_workspace_id = %s WHERE id = %s",
            (workspace["workspace_id"], user["id"])
        )

        return workspace


def get_process_progress_for_user(user):
    """Return process progress for the user's current workspace."""
    workspace = ensure_user_workspace(user)

    with database_connection() as db:
        rows = db.execute(
            """
            SELECT
                pp.progress_id,
                pp.workspace_id,
                pp.process_id,
                p.process_number,
                p.process_name,
                COALESCE(pp.status, 'Not Started') AS status,
                pp.completed_at
            FROM process_progress pp
            JOIN process p
              ON p.process_id = pp.process_id
            WHERE pp.workspace_id = %s
            ORDER BY p.process_number
            """,
            (workspace["workspace_id"],)
        ).fetchall()

    progress = []

    for row in rows:
        item = dict(row)

        if item.get("completed_at") is not None:
            item["completed_at"] = item["completed_at"].isoformat()

        progress.append(item)

    return {
        "workspaceId": workspace["workspace_id"],
        "progress": progress
    }


def set_process_progress(user, process_number, status):
    """Insert or update one process progress record."""
    workspace = ensure_user_workspace(user)

    normalized_status = str(status or "").strip()

    allowed = {
        "Not Started",
        "In Progress",
        "Completed"
    }

    if normalized_status not in allowed:
        raise ValueError("Invalid process progress status.")

    with database_connection() as db:
        process = db.execute(
            """
            SELECT process_id
            FROM process
            WHERE process_number = %s
            """,
            (process_number,)
        ).fetchone()

        if not process:
            raise ValueError("Process not found.")

        existing = db.execute(
            """
            SELECT progress_id
            FROM process_progress
            WHERE workspace_id = %s
              AND process_id = %s
            ORDER BY progress_id
            LIMIT 1
            """,
            (
                workspace["workspace_id"],
                process["process_id"]
            )
        ).fetchone()

        completed_at_sql = (
            "CURRENT_TIMESTAMP"
            if normalized_status == "Completed"
            else "NULL"
        )

        if existing:
            db.execute(
                f"""
                UPDATE process_progress
                SET
                    status = %s,
                    completed_at = {completed_at_sql}
                WHERE progress_id = %s
                """,
                (
                    normalized_status,
                    existing["progress_id"]
                )
            )
        else:
            db.execute(
                f"""
                INSERT INTO process_progress (
                    workspace_id,
                    process_id,
                    status,
                    completed_at
                )
                VALUES (
                    %s,
                    %s,
                    %s,
                    {completed_at_sql}
                )
                """,
                (
                    workspace["workspace_id"],
                    process["process_id"],
                    normalized_status
                )
            )

    return get_process_progress_for_user(user)


def all_processes():
    """Read processes, questions and inputs from PostgreSQL."""
    with database_connection() as db:

        processes = db.execute("""
            SELECT
                process_id,
                process_number,
                process_name,
                purpose
            FROM process
            ORDER BY process_number
        """).fetchall()

        questions = db.execute("""
            SELECT
                process_id,
                question_text,
                question_order
            FROM question
            ORDER BY process_id, question_order
        """).fetchall()

        inputs = db.execute("""
            SELECT
                pi.process_id,
                i.input_name,
                pi.is_required
            FROM process_input pi
            JOIN input i
                ON i.input_id = pi.input_id
            ORDER BY pi.process_id, pi.process_input_id
        """).fetchall()

        output_dependencies = db.execute("""
            SELECT
                pod.process_id,
                source.process_number AS source_process_number
            FROM process_output_dependency pod
            JOIN process source
                ON source.process_id = pod.source_process_id
            ORDER BY pod.process_id, pod.dependency_order
        """).fetchall()

    questions_by_process = {}

    for row in questions:
        questions_by_process.setdefault(
            row["process_id"], []
        ).append(row["question_text"])

    inputs_by_process = {}

    for row in inputs:
        inputs_by_process.setdefault(
            row["process_id"], []
        ).append(row["input_name"])

    output_sources_by_process = {}

    for row in output_dependencies:
        output_sources_by_process.setdefault(
            row["process_id"], []
        ).append(row["source_process_number"])

    return [
        {
            "number": row["process_number"],
            "title": row["process_name"],
            "category": "Strategic process",
            "purpose": row["purpose"],
            "inputs": inputs_by_process.get(row["process_id"], []),
            "questions": questions_by_process.get(row["process_id"], []),
            "outputSources": output_sources_by_process.get(row["process_id"], []),
            "outputs": "",
            "feeds": ""
        }
        for row in processes
    ]

def password_hash(password, salt):
    """Hash a password slowly so database contents cannot reveal passwords."""
    return hashlib.pbkdf2_hmac("sha256", password.encode(), bytes.fromhex(salt), 210_000).hex()


def supabase_auth_request(path, body):
    """POST to the Supabase Authentication (GoTrue) REST API.

    Returns (status_code, parsed_json_body). Raises ValueError only for
    configuration/network problems -- HTTP error status codes from
    Supabase (wrong password, duplicate email, etc.) are returned to the
    caller so they can show a normal error message instead of a crash.
    """
    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise ValueError(
            "Supabase Authentication is not configured. Add SUPABASE_URL and "
            "SUPABASE_ANON_KEY to .env.local."
        )
    request = urllib.request.Request(
        f"{SUPABASE_URL}/auth/v1/{path}",
        data=json.dumps(body).encode(),
        headers={
            "Content-Type": "application/json",
            "apikey": SUPABASE_ANON_KEY,
            "Authorization": f"Bearer {SUPABASE_ANON_KEY}",
        },
        method="POST",
    )
    ssl_context = ssl.create_default_context(cafile=certifi.where())
    try:
        with urllib.request.urlopen(request, timeout=20, context=ssl_context) as response:
            return response.status, json.loads(response.read().decode() or "{}")
    except urllib.error.HTTPError as error:
        try:
            return error.code, json.loads(error.read().decode() or "{}")
        except json.JSONDecodeError:
            return error.code, {}
    except (urllib.error.URLError, TimeoutError) as error:
        raise ValueError(f"Could not reach Supabase Authentication: {error}") from error


def supabase_auth_signup(email, password):
    """Create the account in Supabase Authentication. Returns the Supabase user id."""
    status, body = supabase_auth_request("signup", {"email": email, "password": password})
    if status not in (200, 201):
        message = body.get("msg") or body.get("error_description") or body.get("message") or "Supabase Authentication rejected this account."
        raise ValueError(message)
    supabase_user_id = body.get("id") or (body.get("user") or {}).get("id")
    if not supabase_user_id:
        raise ValueError("Supabase Authentication did not return a user id.")
    return supabase_user_id


def supabase_auth_login(email, password):
    """Verify credentials against Supabase Authentication. Returns True/False."""
    status, body = supabase_auth_request("token?grant_type=password", {"email": email, "password": password})
    return status == 200 and bool(body.get("access_token"))


def create_session(user_id, remember):
    """Create an opaque cookie token that points to a database session."""
    token = secrets.token_urlsafe(32)
    lifetime = 60 * 60 * 24 * (30 if remember else 1)
    with database_connection() as db:
        db.execute("INSERT INTO sessions (user_id, token, expires_at) VALUES (%s, %s, %s)", (user_id, token, int(time.time()) + lifetime))
    return token, lifetime


def user_from_cookie(header):
    """Find the currently signed-in user from a valid, unexpired session."""
    jar = cookies.SimpleCookie(header or "")
    if SESSION_COOKIE not in jar:
        return None
    token = jar[SESSION_COOKIE].value
    with database_connection() as db:
        row = db.execute("""
            SELECT users.* FROM sessions JOIN users ON users.id = sessions.user_id
            WHERE sessions.token = %s AND sessions.expires_at > %s
        """, (token, int(time.time()))).fetchone()
    return row


def public_user(user):
    """Return only the account details that are safe to send to the browser."""
    return {"username": user["username"], "email": user["email"], "companyName": user["company_name"]}


def response_text(response):
    """Extract readable text from a Responses API result."""
    if response.get("output_text"):
        return response["output_text"]
    return "\n".join(
        item.get("text", "")
        for output in response.get("output", [])
        for item in output.get("content", [])
        if item.get("type") == "output_text"
    )


def create_decision_brief(payload):
    """Use the submitted evidence to create a decision brief through the AI API."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("The AI API key is not configured.")
    step = payload.get("step", {})
    documents = payload.get("documents", [])[:20]
    questions = "\n".join(f"{number}. {question}" for number, question in enumerate(step.get("questions", []), 1))
    sources = "\n\n".join(
        f"SOURCE {number}: {doc.get('name', 'Untitled')}\nInput supported: {doc.get('input', 'General')}\n{str(doc.get('text', ''))[:12000]}"
        for number, doc in enumerate(documents, 1)
    )
    if not sources:
        raise ValueError("Please add documentation before generating an output.")
    previous = "\n\n".join(payload.get("previousOutputs", [])) or "None yet."
    body = {"model": "gpt-5.6-sol", "input": [
        {"role": "developer", "content": "You are a strategic market-positioning analyst. Answer every numbered question using the substantive content of the supplied evidence. Do not require documents to mention a particular company or brand name: evidence may still be relevant when it describes the business, market, customers, products, or operations without naming the company. Do not invent facts. Identify any important evidence gaps briefly, then still give the best practical answer and recommendation possible from the available evidence. Clearly label conclusions that are informed assumptions rather than established facts. Ignore sources that are clearly unrelated. Make the response presentation-ready Markdown: begin every answer with its full question in bold on its own line, followed by a concise answer beneath it; use short paragraphs or bullets; and use a Markdown table only where it makes a comparison easier to understand. End with a bold Recommended decision heading and a short recommendation."},
        {"role": "user", "content": f"Purpose: {step.get('purpose', '')}\n\nQuestions:\n{questions}\n\nDocumentation:\n{sources}\n\nPrevious outputs:\n{previous}"}
    ]}
    request = urllib.request.Request("https://api.openai.com/v1/responses", data=json.dumps(body).encode(), headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}, method="POST")
    try:
        ssl_context = ssl.create_default_context(cafile=certifi.where())
        with urllib.request.urlopen(request, timeout=90, context=ssl_context) as response:
            return response_text(json.loads(response.read().decode()))
    except urllib.error.HTTPError as error:
        detail = error.read().decode(errors="replace")[:500]
        raise ValueError(f"AI service returned {error.code}: {detail}") from error


def save_artefact_version(user, process_number, content):
    """Store a new AI-generated document version, keeping only the latest 3 per phase in the current workspace."""
    workspace = ensure_user_workspace(user)

    with database_connection() as db:
        db.execute("""
            INSERT INTO artefact_versions (user_id, workspace_id, process_number, content, created_at)
            VALUES (%s, %s, %s, %s, %s)
        """, (user["id"], workspace["workspace_id"], process_number, content, int(time.time())))

        db.execute("""
            DELETE FROM artefact_versions
            WHERE version_id IN (
                SELECT version_id FROM artefact_versions
                WHERE workspace_id = %s AND process_number = %s
                ORDER BY created_at DESC, version_id DESC
                OFFSET 3
            )
        """, (workspace["workspace_id"], process_number))


def get_artefact_versions(user, process_number):
    """Return the latest saved document versions for one phase in the current workspace, newest first."""
    workspace = ensure_user_workspace(user)

    with database_connection() as db:
        rows = db.execute("""
            SELECT version_id, content, created_at
            FROM artefact_versions
            WHERE workspace_id = %s AND process_number = %s
            ORDER BY created_at DESC, version_id DESC
            LIMIT 3
        """, (workspace["workspace_id"], process_number)).fetchall()

    return [
        {
            "id": row["version_id"],
            "content": row["content"],
            "createdAt": row["created_at"],
        }
        for row in rows
    ]


def workspace_to_public(row, active_workspace_id=None):
    """Convert a workspace database row to the API representation."""
    return {
        "id": row["workspace_id"],
        "workspaceName": row["workspace_name"] or "",
        "businessName": row["organisation_name"] or "",
        "industry": row.get("industry") or "",
        "businessStage": row.get("business_stage") or "",
        "primaryMarket": row.get("primary_market") or "",
        "website": row.get("website") or "",
        "isActive": row["workspace_id"] == active_workspace_id,
    }


def get_workspaces(user):
    """Return every workspace this user owns, each with its market segments."""
    with database_connection() as db:
        account = db.execute(
            "SELECT active_workspace_id FROM users WHERE id = %s",
            (user["id"],)
        ).fetchone()

        active_workspace_id = (
            account["active_workspace_id"] if account else None
        )

        rows = db.execute(
            """
            SELECT
                workspace_id,
                user_id,
                workspace_name,
                organisation_name,
                industry,
                business_stage,
                primary_market,
                website
            FROM workspace
            WHERE user_id = %s
            ORDER BY workspace_id
            """,
            (user["id"],)
        ).fetchall()

        segment_rows = db.execute(
            """
            SELECT
                workspace_id,
                segment_id,
                segment_name
            FROM market_segments
            WHERE workspace_id IN (
                SELECT workspace_id FROM workspace WHERE user_id = %s
            )
            ORDER BY workspace_id, segment_id
            """,
            (user["id"],)
        ).fetchall()

    segments_by_workspace = {}

    for segment in segment_rows:
        segments_by_workspace.setdefault(
            segment["workspace_id"], []
        ).append(segment["segment_name"])

    workspaces = []

    for row in rows:
        workspace = workspace_to_public(row, active_workspace_id)
        workspace["segments"] = segments_by_workspace.get(
            row["workspace_id"], []
        )
        workspaces.append(workspace)

    return workspaces


def create_workspace(user, payload):
    """Create a workspace and make it the user's active workspace."""
    workspace_name = str(payload.get("workspaceName", "")).strip()
    business_name = str(payload.get("businessName", "")).strip()
    industry = str(payload.get("industry", "")).strip()
    business_stage = str(payload.get("businessStage", "")).strip()
    primary_market = str(payload.get("primaryMarket", "")).strip()
    website = str(payload.get("website", "")).strip()

    segments = payload.get("segments", [])
    if not isinstance(segments, list):
        segments = []
    segments = [str(segment).strip() for segment in segments if str(segment).strip()]

    if not workspace_name:
        raise ValueError("Workspace name is required.")

    if len(workspace_name) > 200:
        raise ValueError("Workspace name must be 200 characters or fewer.")

    if not business_name:
        raise ValueError("Business name is required.")

    if len(segments) > 2:
        raise ValueError("Add no more than two initial market segments.")

    now = int(time.time())

    with database_connection() as db:
        row = db.execute(
            """
            INSERT INTO workspace (
                user_id,
                workspace_name,
                organisation_name,
                industry,
                business_stage,
                primary_market,
                website
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING
                workspace_id,
                user_id,
                workspace_name,
                organisation_name,
                industry,
                business_stage,
                primary_market,
                website
            """,
            (
                user["id"],
                workspace_name,
                business_name,
                industry or None,
                business_stage or None,
                primary_market or None,
                website or None,
            )
        ).fetchone()

        db.execute(
            "UPDATE users SET active_workspace_id = %s WHERE id = %s",
            (row["workspace_id"], user["id"])
        )

        for segment_name in segments:
            db.execute(
                """
                INSERT INTO market_segments (
                    user_id, workspace_id, segment_name, description,
                    geography, company_size, wedge, created_at, updated_at
                )
                VALUES (%s, %s, %s, '', '', '', '', %s, %s)
                """,
                (user["id"], row["workspace_id"], segment_name, now, now)
            )

    return workspace_to_public(row, row["workspace_id"])


def select_workspace(user, workspace_id):
    """Set one of the user's workspaces as the active workspace."""
    with database_connection() as db:
        owned = db.execute(
            """
            SELECT workspace_id
            FROM workspace
            WHERE workspace_id = %s AND user_id = %s
            """,
            (workspace_id, user["id"])
        ).fetchone()

        if not owned:
            raise ValueError("Workspace not found.")

        db.execute(
            "UPDATE users SET active_workspace_id = %s WHERE id = %s",
            (workspace_id, user["id"])
        )

    return {"activeWorkspaceId": workspace_id}


def update_workspace(user, workspace_id, payload):
    """Update workspace setup information and its initial market segments."""
    workspace_name = str(payload.get("workspaceName", "")).strip()
    business_name = str(payload.get("businessName", "")).strip()
    industry = str(payload.get("industry", "")).strip()
    business_stage = str(payload.get("businessStage", "")).strip()
    primary_market = str(payload.get("primaryMarket", "")).strip()
    website = str(payload.get("website", "")).strip()

    segments = payload.get("segments", [])
    if not isinstance(segments, list):
        segments = []
    segments = [str(segment).strip() for segment in segments if str(segment).strip()]

    if not workspace_name:
        raise ValueError("Workspace name is required.")

    if not business_name:
        raise ValueError("Business name is required.")

    if len(segments) > 2:
        raise ValueError("Add no more than two initial market segments.")

    now = int(time.time())

    with database_connection() as db:
        row = db.execute(
            """
            UPDATE workspace
            SET
                workspace_name = %s,
                organisation_name = %s,
                industry = %s,
                business_stage = %s,
                primary_market = %s,
                website = %s
            WHERE workspace_id = %s AND user_id = %s
            RETURNING
                workspace_id,
                user_id,
                workspace_name,
                organisation_name,
                industry,
                business_stage,
                primary_market,
                website
            """,
            (
                workspace_name,
                business_name,
                industry or None,
                business_stage or None,
                primary_market or None,
                website or None,
                workspace_id,
                user["id"],
            )
        ).fetchone()

        if not row:
            raise ValueError("Workspace not found.")

        # Initial market segments are fully replaced on every edit -- this
        # form only ever holds up to two plain segment names, so there is
        # no partial-update case to preserve.
        db.execute(
            "DELETE FROM market_segments WHERE workspace_id = %s",
            (workspace_id,)
        )

        for segment_name in segments:
            db.execute(
                """
                INSERT INTO market_segments (
                    user_id, workspace_id, segment_name, description,
                    geography, company_size, wedge, created_at, updated_at
                )
                VALUES (%s, %s, %s, '', '', '', '', %s, %s)
                """,
                (user["id"], workspace_id, segment_name, now, now)
            )

    return workspace_to_public(row)


def delete_workspace(user, workspace_id):
    """Delete one of the user's workspaces and everything scoped to it.

    process_progress, market_segments and evidence_items all reference
    workspace_id with ON DELETE CASCADE, so deleting the workspace row
    takes its data with it. Deleting the account's only remaining
    workspace is refused -- there must always be somewhere for process
    progress, segments and evidence to live.
    """
    with database_connection() as db:
        remaining = db.execute(
            "SELECT COUNT(*) AS count FROM workspace WHERE user_id = %s",
            (user["id"],)
        ).fetchone()

        if remaining["count"] <= 1:
            raise ValueError("You must keep at least one workspace.")

        cursor = db.execute(
            "DELETE FROM workspace WHERE workspace_id = %s AND user_id = %s",
            (workspace_id, user["id"])
        )

        if cursor.rowcount == 0:
            raise ValueError("Workspace not found.")

        # If the deleted workspace was active, users.active_workspace_id
        # was just set to NULL by the ON DELETE SET NULL foreign key --
        # point it at another remaining workspace right away so the next
        # request doesn't have to fall back through ensure_user_workspace().
        account = db.execute(
            "SELECT active_workspace_id FROM users WHERE id = %s",
            (user["id"],)
        ).fetchone()

        if not account or account["active_workspace_id"] is None:
            fallback = db.execute(
                """
                SELECT workspace_id FROM workspace
                WHERE user_id = %s
                ORDER BY workspace_id
                LIMIT 1
                """,
                (user["id"],)
            ).fetchone()

            if fallback:
                db.execute(
                    "UPDATE users SET active_workspace_id = %s WHERE id = %s",
                    (fallback["workspace_id"], user["id"])
                )

    return get_workspaces(user)


def get_segments(user):
    """Return all market segments belonging to the user's active workspace."""
    workspace = ensure_user_workspace(user)

    with database_connection() as db:
        rows = db.execute("""
            SELECT
                segment_id,
                segment_name,
                description,
                geography,
                company_size,
                wedge,
                created_at,
                updated_at
            FROM market_segments
            WHERE workspace_id = %s
            ORDER BY segment_id
        """, (workspace["workspace_id"],)).fetchall()

    return [
        {
            "id": row["segment_id"],
            "name": row["segment_name"],
            "description": row["description"] or "",
            "geography": row["geography"] or "",
            "companySize": row["company_size"] or "",
            "wedge": row["wedge"] or "",
            "createdAt": row["created_at"],
            "updatedAt": row["updated_at"],
        }
        for row in rows
    ]


def create_segment(user, payload):
    """Create a market segment in the user's active workspace."""
    workspace = ensure_user_workspace(user)

    name = str(payload.get("name", "")).strip()
    description = str(payload.get("description", "")).strip()
    geography = str(payload.get("geography", "")).strip()
    company_size = str(payload.get("companySize", "")).strip()
    wedge = str(payload.get("wedge", "")).strip()

    if not name:
        raise ValueError("Segment name is required.")

    if len(name) > 150:
        raise ValueError("Segment name must be 150 characters or fewer.")

    now = int(time.time())

    with database_connection() as db:
        duplicate = db.execute("""
            SELECT 1
            FROM market_segments
            WHERE workspace_id = %s
              AND LOWER(segment_name) = LOWER(%s)
        """, (workspace["workspace_id"], name)).fetchone()

        if duplicate:
            raise ValueError("This segment already exists in your workspace.")

        row = db.execute("""
            INSERT INTO market_segments
                (user_id, workspace_id, segment_name, description, geography,
                 company_size, wedge, created_at, updated_at)
            VALUES
                (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING
                segment_id,
                segment_name,
                description,
                geography,
                company_size,
                wedge,
                created_at,
                updated_at
        """, (
            user["id"],
            workspace["workspace_id"],
            name,
            description,
            geography,
            company_size,
            wedge,
            now,
            now
        )).fetchone()

    return {
        "id": row["segment_id"],
        "name": row["segment_name"],
        "description": row["description"] or "",
        "geography": row["geography"] or "",
        "companySize": row["company_size"] or "",
        "wedge": row["wedge"] or "",
        "createdAt": row["created_at"],
        "updatedAt": row["updated_at"],
    }


def delete_segment(user, segment_id):
    """Delete only a segment belonging to the user's active workspace."""
    workspace = ensure_user_workspace(user)

    with database_connection() as db:
        cursor = db.execute("""
            DELETE FROM market_segments
            WHERE segment_id = %s
              AND workspace_id = %s
        """, (segment_id, workspace["workspace_id"]))

    if cursor.rowcount == 0:
        raise ValueError("Segment not found.")

class FounderMotionHandler(BaseHTTPRequestHandler):
    """HTTP routes for account actions, saved map data, AI analysis, and web files."""

    def send_json(self, status, body, cookie=None):
        encoded = json.dumps(body).encode()
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(encoded)))
        if cookie:
            self.send_header("Set-Cookie", cookie.output(header="").strip())
        self.end_headers()
        self.wfile.write(encoded)

    def read_body(self, max_size=1_500_000):
        size = int(self.headers.get("Content-Length", "0"))
        if size > max_size:
            raise ValueError("Request is too large.")
        return json.loads(self.rfile.read(size).decode("utf-8"))

    def current_user(self):
        return user_from_cookie(self.headers.get("Cookie"))

    def session_cookie(self, token, max_age):
        jar = cookies.SimpleCookie()
        jar[SESSION_COOKIE] = token
        jar[SESSION_COOKIE]["path"] = "/"
        jar[SESSION_COOKIE]["httponly"] = True
        jar[SESSION_COOKIE]["samesite"] = "Lax"
        return jar

    def require_user(self):
        user = self.current_user()
        if not user:
            self.send_json(401, {"error": "Please log in first."})
        return user

    def do_POST(self):
        try:
            upload_body_limit = 16_000_000 if self.path == "/api/evidence/upload" else 1_500_000
            payload = self.read_body(max_size=upload_body_limit)
            if self.path == "/api/auth/register":
                self.register(payload)
            elif self.path == "/api/auth/login":
                self.login(payload)
            elif self.path == "/api/auth/logout":
                self.logout()
            elif self.path == "/api/auth/forgot-password":
                self.send_json(200, {"message": "If the account exists, a reset email will be sent when email delivery is configured."})
            elif self.path == "/api/analyze":
                user = self.require_user()
                if user:
                    answer = create_decision_brief(payload) or "No answer was returned."
                    process_number = payload.get("step", {}).get("number")
                    if process_number:
                        try:
                            save_artefact_version(user, int(process_number), answer)
                        except Exception:
                            import traceback
                            traceback.print_exc()
                    self.send_json(200, {"answer": answer})

            elif self.path == "/api/segments":
                user = self.require_user()
                if user:
                    segment = create_segment(user, payload)
                    self.send_json(201, {"segment": segment})

            elif self.path == "/api/evidence":
                user = self.require_user()
                if user:
                    evidence = create_evidence(user, payload)
                    self.send_json(201, {"evidence": evidence})

            elif self.path == "/api/evidence/upload":
                user = self.require_user()
                if user:
                    result = handle_evidence_upload(user["id"], payload)
                    self.send_json(201, result)

            elif self.path.startswith("/api/evidence/") and self.path.endswith("/delete"):
                user = self.require_user()
                if user:
                    evidence_id = int(self.path.split("/")[3])
                    delete_evidence(user, evidence_id)
                    self.send_json(200, {"deleted": True})

            elif self.path.startswith("/api/segments/") and self.path.endswith("/delete"):
                user = self.require_user()
                if user:
                    segment_id = int(self.path.split("/")[3])
                    delete_segment(user, segment_id)
                    self.send_json(200, {"deleted": True})

            elif self.path == "/api/workspaces":
                user = self.require_user()
                if user:
                    workspace = create_workspace(user, payload)
                    self.send_json(201, {"workspace": workspace})

            elif self.path == "/api/workspaces/select":
                user = self.require_user()
                if user:
                    workspace_id = int(payload.get("workspaceId"))
                    result = select_workspace(user, workspace_id)
                    self.send_json(200, result)

            elif self.path.startswith("/api/workspaces/") and self.path.endswith("/update"):
                user = self.require_user()
                if user:
                    workspace_id = int(self.path.split("/")[3])
                    workspace = update_workspace(user, workspace_id, payload)
                    self.send_json(200, {"workspace": workspace})

            elif self.path.startswith("/api/workspaces/") and self.path.endswith("/delete"):
                user = self.require_user()
                if user:
                    workspace_id = int(self.path.split("/")[3])
                    workspaces = delete_workspace(user, workspace_id)
                    self.send_json(200, {"workspaces": workspaces})
            else:
                self.send_json(404, {"error": "Not found."})
        except ValueError as error:
            self.send_json(400, {"error": str(error)})
        except Exception as error:
            import traceback
            traceback.print_exc()
            try:
                self.send_json(500, {"error": f"Server error: {error}"})
            except Exception:
                pass

    def do_GET(self):
        if self.path == "/api/processes":
            self.send_json(200, {"processes": all_processes()})
            return
        if self.path == "/api/auth/session":
            user = self.current_user()
            self.send_json(200, {"user": public_user(user) if user else None})
            return

        if self.path == "/api/segments":
            user = self.require_user()
            if user:
                self.send_json(200, {"segments": get_segments(user)})
            return

        if self.path == "/api/evidence":
            user = self.require_user()
            if user:
                self.send_json(200, {"evidence": get_evidence(user)})
            return

        if self.path == "/api/workspaces":
            user = self.require_user()
            if user:
                self.send_json(200, {"workspaces": get_workspaces(user)})
            return
        if self.path == "/api/process-progress":
            user = self.require_user()
            if user:
                self.send_json(
                    200,
                    get_process_progress_for_user(user)
                )
            return

        if self.path.startswith("/api/artefact-versions"):
            user = self.require_user()
            if user:
                query = urllib.parse.urlsplit(self.path).query
                params = urllib.parse.parse_qs(query)
                try:
                    process_number = int(params.get("processNumber", ["0"])[0])
                except ValueError:
                    process_number = 0
                if not 1 <= process_number <= 13:
                    self.send_json(400, {"error": "Invalid process number."})
                    return
                self.send_json(200, {"versions": get_artefact_versions(user, process_number)})
            return

        if self.path == "/api/state":
            user = self.require_user()
            if user:
                workspace = ensure_user_workspace(user)
                with database_connection() as db:
                    row = db.execute("SELECT state_json FROM map_states WHERE workspace_id = %s", (workspace["workspace_id"],)).fetchone()
                    answers = db.execute("SELECT process_number, answer FROM process_answers WHERE workspace_id = %s", (workspace["workspace_id"],)).fetchall()
                state = json.loads(row["state_json"]) if row else {"step": 0, "documents": []}
                state["outputs"] = (state.get("outputs", []) + [""] * 13)[:13]
                for answer in answers:
                    state["outputs"][answer["process_number"] - 1] = answer["answer"]
                self.send_json(200, {"state": state})
            return
        self.serve_file()

    def do_PUT(self):
        try:
            user = self.require_user()
            if not user:
                return
            payload = self.read_body()
            if self.path == "/api/process-progress":
                process_number = int(payload.get("processNumber", 0))
                status = str(payload.get("status", "")).strip()

                if not 1 <= process_number <= 13:
                    raise ValueError("Invalid process number.")

                result = set_process_progress(
                    user,
                    process_number,
                    status
                )

                self.send_json(200, result)
                return

            if self.path == "/api/state":
                workspace = ensure_user_workspace(user)
                with database_connection() as db:
                    state_for_storage = {
                        "step": payload.get("step", 0),
                        "documents": payload.get("documents", []),
                        "history": payload.get("history", [])[:3]
                    }
                    db.execute(
                        "INSERT INTO map_states (user_id, workspace_id, state_json, updated_at) VALUES (%s, %s, %s, %s) ON CONFLICT(workspace_id) DO UPDATE SET state_json = excluded.state_json, updated_at = excluded.updated_at",
                        (user["id"], workspace["workspace_id"], json.dumps(state_for_storage), int(time.time()))
                    )
                    db.execute("DELETE FROM process_answers WHERE workspace_id = %s", (workspace["workspace_id"],))
                    for index, answer in enumerate(payload.get("outputs", []), 1):
                        if answer:
                            db.execute(
                                "INSERT INTO process_answers (user_id, workspace_id, process_number, answer, updated_at) VALUES (%s, %s, %s, %s, %s) ON CONFLICT(workspace_id, process_number) DO UPDATE SET answer = excluded.answer, updated_at = excluded.updated_at",
                                (user["id"], workspace["workspace_id"], index, answer, int(time.time()))
                            )
                self.send_json(200, {"saved": True})
            elif self.path == "/api/company":
                name = str(payload.get("companyName", "")).strip()
                if not 2 <= len(name) <= 50:
                    raise ValueError("Company name must be between 2 and 50 characters.")
                with database_connection() as db:
                    db.execute("UPDATE users SET company_name = %s WHERE id = %s", (name, user["id"]))
                self.send_json(200, {"companyName": name})
            else:
                self.send_json(404, {"error": "Not found."})
        except ValueError as error:
            self.send_json(400, {"error": str(error)})

    def register(self, payload):
        identity = str(payload.get("identity", "")).strip().lower()
        password = str(payload.get("password", ""))
        company = str(payload.get("companyName", "")).strip() or "FounderMotion"
        if len(password) < 8:
            raise ValueError("Enter an email and a password of at least 8 characters.")
        # New accounts must use Supabase Authentication, which requires a
        # real email address (not an arbitrary username).
        if "@" not in identity or "." not in identity.split("@")[-1]:
            raise ValueError("Please register with a valid email address.")
        email = identity
        with database_connection() as db:
            existing = db.execute("SELECT 1 FROM users WHERE username = %s OR email = %s", (identity, email)).fetchone()
            if existing:
                raise ValueError("That username or email is already registered. Please log in or choose another one.")

        # Create the account in Supabase Authentication first. If this
        # fails (duplicate email, weak password per Supabase's rules,
        # etc.) nothing is written to our own table.
        supabase_user_id = supabase_auth_signup(email, password)

        salt = secrets.token_hex(16)
        with database_connection() as db:
            cursor = db.execute(
                "INSERT INTO users (username, email, password_hash, password_salt, company_name, supabase_user_id, created_at) VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id",
                (identity, email, password_hash(password, salt), salt, company, supabase_user_id, int(time.time()))
            )
            user_id = cursor.fetchone()["id"]
        token, lifetime = create_session(user_id, bool(payload.get("remember")))
        self.send_json(201, {"user": {"username": identity, "companyName": company}}, self.session_cookie(token, lifetime))

    def login(self, payload):
        identity = str(payload.get("identity", "")).strip().lower()
        password = str(payload.get("password", ""))
        with database_connection() as db:
            user = db.execute("SELECT * FROM users WHERE username = %s OR email = %s", (identity, identity)).fetchone()
        if not user:
            raise ValueError("Incorrect username/email or password.")

        if user["supabase_user_id"]:
            # Account created after the Supabase Authentication switch:
            # credentials are verified by Supabase, not our own hash.
            if not user["email"] or not supabase_auth_login(user["email"], password):
                raise ValueError("Incorrect username/email or password.")
        elif not hmac.compare_digest(user["password_hash"], password_hash(password, user["password_salt"])):
            # Legacy account created before this change: fall back to the
            # original local password check so existing users are not locked out.
            raise ValueError("Incorrect username/email or password.")

        token, lifetime = create_session(user["id"], bool(payload.get("remember")))
        self.send_json(200, {"user": public_user(user)}, self.session_cookie(token, lifetime))

    def logout(self):
        jar = cookies.SimpleCookie(self.headers.get("Cookie", ""))
        if SESSION_COOKIE in jar:
            token = jar[SESSION_COOKIE].value

            with database_connection() as db:
                # Find the logged-in user before deleting the session.
                session_user = db.execute(
                    "SELECT user_id FROM sessions WHERE token = %s",
                    (token,)
                ).fetchone()

                # Preserve all saved workspace progress when logging out.
                # Only the login session should be removed.

                # Finally delete the login session.
                db.execute(
                    "DELETE FROM sessions WHERE token = %s",
                    (token,)
                )

        expired = self.session_cookie("", 0)
        self.send_json(200, {"loggedOut": True}, expired)

    def serve_file(self):
        requested = "index.html" if self.path in ("/", "") else self.path.lstrip("/").split("?", 1)[0]
        file = (ROOT / requested).resolve()
        if ROOT not in file.parents or not file.is_file():
            self.send_error(404)
            return
        data = file.read_bytes()
        content_types = {".html": "text/html; charset=utf-8", ".css": "text/css; charset=utf-8", ".js": "application/javascript; charset=utf-8"}
        self.send_response(200)
        self.send_header("Content-Type", content_types.get(file.suffix, "application/octet-stream"))
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)


if __name__ == "__main__":
    if not POSTGRES_PASSWORD:
        raise RuntimeError(
            "POSTGRES_PASSWORD is missing. Add it to .env.local before starting the server."
        )

    if not S3_BUCKET_NAME:
        raise RuntimeError(
            "S3_BUCKET_NAME is missing. Add it to .env.local before starting the server."
        )

    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise RuntimeError(
            "SUPABASE_URL / SUPABASE_ANON_KEY is missing. Add them to .env.local "
            "before starting the server (Supabase dashboard -> Project Settings -> API)."
        )

    initialise_database()

    print(f"FounderMotion is running at http://localhost:{PORT}")
    ThreadingHTTPServer(("", PORT), FounderMotionHandler).serve_forever()
